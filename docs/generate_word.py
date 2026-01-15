#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ERP系统用户操作手册Word文档生成脚本
运行方式: python generate_word.py
输出文件: output/ERP系统用户操作手册.docx

截图说明:
将系统截图放入 screenshots/ 目录，按以下命名规则：
- login.png           登录页面
- dashboard.png       工作台
- customer_list.png   客户列表
- customer_form.png   客户表单
- project_list.png    项目列表
- project_form.png    项目表单
- bom_list.png        BOM列表
- purchase_request.png 采购申请
- purchase_order.png  采购订单
- sales_order.png     销售订单
- inventory.png       库存查询
- workflow.png        审批流程
- ... (更多截图)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


class ManualGenerator:
    def __init__(self):
        self.doc = Document()
        self.screenshot_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "screenshots")
        self._setup_styles()
    
    def _setup_styles(self):
        """设置文档样式"""
        # 设置正文样式
        style = self.doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 设置标题样式
        for i in range(1, 5):
            style_name = f'Heading {i}'
            if style_name in self.doc.styles:
                style = self.doc.styles[style_name]
                style.font.name = '微软雅黑'
                style.font.bold = True
                style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                if i == 1:
                    style.font.size = Pt(22)
                    style.font.color.rgb = RGBColor(0x1E, 0x88, 0xE5)
                elif i == 2:
                    style.font.size = Pt(16)
                    style.font.color.rgb = RGBColor(0x42, 0x42, 0x42)
                elif i == 3:
                    style.font.size = Pt(14)
                    style.font.color.rgb = RGBColor(0x42, 0x42, 0x42)
                else:
                    style.font.size = Pt(12)
    
    def add_cover(self, title, subtitle, version, date):
        """添加封面"""
        # 空几行
        for _ in range(6):
            self.doc.add_paragraph()
        
        # 标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1E, 0x88, 0xE5)
        
        # 副标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
        
        # 空几行
        for _ in range(8):
            self.doc.add_paragraph()
        
        # 版本信息
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"版本：{version}")
        run.font.size = Pt(12)
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"日期：{date}")
        run.font.size = Pt(12)
        
        # 分页
        self.doc.add_page_break()
    
    def add_toc(self):
        """添加目录"""
        self.doc.add_heading('目录', level=1)
        
        toc_items = [
            "1. 系统概述",
            "2. 登录与个人中心",
            "3. 工作台",
            "4. CRM客户管理",
            "5. PLM产品研发",
            "6. ERP销售执行",
            "7. ERP采购管理",
            "8. ERP供应商管理",
            "9. ERP库存管理",
            "10. ERP财务管理",
            "11. MES生产执行",
            "12. MES设备管理",
            "13. OA协同办公",
            "14. 系统管理",
            "15. 报表分析",
            "16. 常见问题FAQ",
        ]
        
        for item in toc_items:
            p = self.doc.add_paragraph(item)
            p.paragraph_format.left_indent = Inches(0.5)
        
        self.doc.add_page_break()
    
    def add_heading(self, text, level=1):
        """添加标题"""
        self.doc.add_heading(text, level=level)
    
    def add_paragraph(self, text, bold=False, indent=False):
        """添加段落"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.bold = bold
        if indent:
            p.paragraph_format.left_indent = Inches(0.3)
        return p
    
    def add_bullet_list(self, items, indent=0):
        """添加无序列表"""
        for item in items:
            p = self.doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.3 + indent * 0.3)
    
    def add_numbered_list(self, items):
        """添加有序列表"""
        for item in items:
            p = self.doc.add_paragraph(item, style='List Number')
            p.paragraph_format.left_indent = Inches(0.3)
    
    def add_table(self, headers, rows):
        """添加表格"""
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # 设置表头样式
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            # 设置背景色
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '1E88E5')
            header_cells[i]._tc.get_or_add_tcPr().append(shading)
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
        # 数据行
        for row_idx, row in enumerate(rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, value in enumerate(row):
                row_cells[col_idx].text = str(value)
                for paragraph in row_cells[col_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        
        self.doc.add_paragraph()  # 表格后空行
        return table
    
    def add_image(self, image_name, caption="", width=5.5):
        """添加图片（如果存在）"""
        image_path = os.path.join(self.screenshot_dir, image_name)
        
        if os.path.exists(image_path):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(image_path, width=Inches(width))
            
            if caption:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"图：{caption}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
        else:
            # 添加占位符
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"[ 请添加截图: {image_name} ]")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.italic = True
            
            if caption:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"图：{caption}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
        
        self.doc.add_paragraph()
    
    def add_note(self, text):
        """添加注意事项"""
        p = self.doc.add_paragraph()
        run = p.add_run("💡 提示：")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x43, 0xA0, 0x47)
        run = p.add_run(text)
        run.font.color.rgb = RGBColor(0x43, 0xA0, 0x47)
        p.paragraph_format.left_indent = Inches(0.3)
    
    def add_warning(self, text):
        """添加警告"""
        p = self.doc.add_paragraph()
        run = p.add_run("⚠️ 注意：")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0x98, 0x00)
        run = p.add_run(text)
        run.font.color.rgb = RGBColor(0xFF, 0x98, 0x00)
        p.paragraph_format.left_indent = Inches(0.3)
    
    def add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()
    
    def save(self, filename):
        """保存文档"""
        output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, filename)
        self.doc.save(output_path)
        print(f"✅ Word文档已生成: {output_path}")
        return output_path


def create_manual():
    """创建用户操作手册"""
    gen = ManualGenerator()
    
    # 封面
    gen.add_cover(
        "ERP系统用户操作手册",
        "面向非标/项目型制造企业的一体化管理平台",
        "V1.0",
        "2026年1月"
    )
    
    # 目录
    gen.add_toc()
    
    # ========== 第1章：系统概述 ==========
    gen.add_heading("1. 系统概述", level=1)
    
    gen.add_heading("1.1 系统简介", level=2)
    gen.add_paragraph(
        "本ERP系统是一套面向制造企业的综合管理平台，集成了CRM客户管理、PLM产品研发、"
        "ERP资源管理、MES生产执行、OA协同办公等核心功能模块，帮助企业实现从销售、研发、"
        "采购、生产到财务的全流程数字化管理。"
    )
    
    gen.add_heading("1.2 系统架构", level=2)
    gen.add_paragraph("系统采用模块化设计，各模块之间数据自动流转：")
    gen.add_bullet_list([
        "CRM客户管理：销售线索 → 销售商机 → 客户档案 → 报价 → 合同",
        "PLM产品研发：需求管理 → 方案设计 → 项目管理 → BOM管理 → 图纸管理",
        "ERP销售执行：销售订单 → 发货管理 → 应收账款 → 回款计划",
        "ERP采购管理：MRP计划 → 采购申请 → 采购订单 → 到货质检",
        "MES生产执行：生产计划 → APS排程 → 工单派工 → 质量检验",
    ])
    gen.add_image("system_architecture.png", "系统架构图")
    
    gen.add_heading("1.3 浏览器要求", level=2)
    gen.add_table(
        ["项目", "要求"],
        [
            ["推荐浏览器", "Chrome 90+、Firefox 88+、Edge 90+"],
            ["屏幕分辨率", "建议 1920×1080 或以上"],
            ["网络环境", "内网或VPN访问"],
        ]
    )
    
    gen.add_page_break()
    
    # ========== 第2章：登录与个人中心 ==========
    gen.add_heading("2. 登录与个人中心", level=1)
    
    gen.add_heading("2.1 系统登录", level=2)
    gen.add_paragraph("登录步骤：")
    gen.add_numbered_list([
        "打开浏览器，输入系统地址",
        "在登录页面输入用户名和密码",
        "点击「登录」按钮进入系统",
    ])
    gen.add_image("login.png", "登录页面")
    gen.add_note("默认管理员账号：admin / admin123（首次登录后请立即修改密码）")
    
    gen.add_heading("2.2 个人中心", level=2)
    gen.add_paragraph("点击右上角头像，可进入个人中心：")
    gen.add_table(
        ["功能", "说明"],
        [
            ["个人资料", "查看/修改姓名、邮箱、手机、头像等"],
            ["修改密码", "修改登录密码（需输入原密码）"],
            ["通知设置", "设置接收通知的方式"],
        ]
    )
    gen.add_image("profile.png", "个人中心")
    
    gen.add_page_break()
    
    # ========== 第3章：工作台 ==========
    gen.add_heading("3. 工作台（Dashboard）", level=1)
    
    gen.add_heading("3.1 功能概述", level=2)
    gen.add_paragraph("工作台是登录后的首页，提供：")
    gen.add_bullet_list([
        "待办事项：显示待审批、待处理的任务",
        "快捷入口：常用功能的快速访问",
        "数据概览：关键业务指标的可视化展示",
        "公告通知：系统公告和重要通知",
    ])
    gen.add_image("dashboard.png", "工作台")
    
    gen.add_page_break()
    
    # ========== 第4章：CRM客户管理 ==========
    gen.add_heading("4. CRM客户管理", level=1)
    
    gen.add_heading("4.1 销售线索", level=2)
    gen.add_paragraph("路径：CRM客户管理 → 销售线索")
    gen.add_paragraph("销售线索是潜在客户的初始信息录入。")
    gen.add_table(
        ["操作", "说明"],
        [
            ["新建线索", "录入潜在客户的基本信息"],
            ["转为商机", "将有意向的线索转化为销售商机"],
            ["分配", "将线索分配给指定销售人员跟进"],
        ]
    )
    gen.add_image("leads.png", "销售线索列表")
    
    gen.add_heading("4.2 销售商机", level=2)
    gen.add_paragraph("路径：CRM客户管理 → 销售商机")
    gen.add_paragraph("管理销售机会的跟进过程，商机阶段包括：")
    gen.add_numbered_list([
        "初步接触",
        "需求确认",
        "方案报价",
        "商务谈判",
        "签约成交 / 丢单",
    ])
    gen.add_image("opportunities.png", "销售商机")
    
    gen.add_heading("4.3 客户档案", level=2)
    gen.add_paragraph("路径：CRM客户管理 → 客户档案")
    gen.add_paragraph("管理正式客户的详细信息，包括：")
    gen.add_bullet_list([
        "客户基本信息（公司名称、地址、行业、规模）",
        "联系人管理（支持多联系人）",
        "银行账户信息",
        "关联项目、订单、合同查看",
    ])
    gen.add_image("customer_list.png", "客户列表")
    gen.add_image("customer_form.png", "客户表单")
    
    gen.add_heading("4.4 客户跟进", level=2)
    gen.add_paragraph("路径：CRM客户管理 → 客户跟进")
    gen.add_paragraph("记录与客户的沟通历史：")
    gen.add_table(
        ["字段", "说明"],
        [
            ["跟进方式", "电话、拜访、邮件、微信等"],
            ["跟进内容", "沟通的具体内容"],
            ["下次跟进计划", "计划的下次跟进时间和内容"],
        ]
    )
    
    gen.add_heading("4.5 销售报价", level=2)
    gen.add_paragraph("路径：CRM客户管理 → 销售报价")
    gen.add_paragraph("创建和管理销售报价单，操作流程：")
    gen.add_numbered_list([
        "点击「新建报价」",
        "选择客户",
        "添加产品明细（产品、数量、单价、折扣）",
        "设置有效期",
        "保存/提交审批",
        "审批通过后可「转为订单」",
    ])
    gen.add_image("quotation.png", "销售报价")
    
    gen.add_page_break()
    
    # ========== 第5章：PLM产品研发 ==========
    gen.add_heading("5. PLM产品研发", level=1)
    
    gen.add_heading("5.1 项目列表", level=2)
    gen.add_paragraph("路径：PLM产品研发 → 项目列表")
    gen.add_paragraph("管理所有项目的生命周期。")
    gen.add_image("project_list.png", "项目列表")
    
    gen.add_heading("5.1.1 新建项目", level=3)
    gen.add_numbered_list([
        "点击「新建项目」",
        "填写项目信息：项目名称、项目编号、客户、关联销售订单、预算金额、计划开始/结束日期",
        "保存",
    ])
    gen.add_image("project_form.png", "项目表单")
    
    gen.add_heading("5.1.2 关联销售订单", level=3)
    gen.add_paragraph("创建项目时可选择关联销售订单：")
    gen.add_bullet_list([
        "选择销售订单后，系统会显示订单详情（产品、金额、交期等）",
        "自动填充客户、预算金额、结束日期",
        "已关联项目的订单不会出现在可选列表中",
        "项目取消后，关联的订单会被释放",
    ])
    gen.add_note("此功能帮助项目经理快速了解项目背景，准确填写项目信息。")
    
    gen.add_heading("5.2 任务管理", level=2)
    gen.add_paragraph("路径：PLM产品研发 → 任务管理")
    gen.add_paragraph("管理项目下的具体任务：")
    gen.add_table(
        ["字段", "说明"],
        [
            ["任务名称", "任务的名称"],
            ["所属项目", "任务所属的项目"],
            ["负责人", "任务的负责人"],
            ["计划时间", "计划开始/结束时间"],
            ["优先级", "高/中/低"],
            ["任务状态", "待开始/进行中/已完成/已取消"],
        ]
    )
    gen.add_image("task_list.png", "任务列表")
    
    gen.add_heading("5.3 甘特图", level=2)
    gen.add_paragraph("路径：PLM产品研发 → 甘特图")
    gen.add_paragraph("以甘特图形式展示项目进度：")
    gen.add_bullet_list([
        "支持拖拽调整任务时间",
        "支持任务依赖关系显示",
        "支持按项目/人员筛选",
    ])
    gen.add_image("gantt.png", "甘特图")
    
    gen.add_heading("5.4 BOM管理", level=2)
    gen.add_paragraph("路径：PLM产品研发 → BOM管理")
    gen.add_paragraph("管理项目的物料清单（Bill of Materials）。")
    gen.add_image("bom_list.png", "BOM列表")
    
    gen.add_heading("5.4.1 BOM采购状态", level=3)
    gen.add_paragraph("BOM物料的采购状态会实时更新：")
    gen.add_table(
        ["状态", "说明"],
        [
            ["未下单", "尚未生成采购申请"],
            ["申请中", "已生成采购申请，待审批"],
            ["已批准", "采购申请已审批通过"],
            ["已下单", "已生成采购订单并确认"],
            ["在途中", "供应商已发货"],
            ["已收货", "已完成到货验收"],
            ["部分收货", "部分数量已收货"],
            ["已退货", "物料已退回供应商"],
            ["已取消", "采购已取消"],
        ]
    )
    
    gen.add_heading("5.4.2 导出询价BOM", level=3)
    gen.add_numbered_list([
        "使用筛选条件（有图/无图、物料类型、版本/品牌）",
        "选择需要的物料",
        "点击「导出筛选结果(询价用)」或「导出选中(询价用)」",
        "将导出的Excel发送给供应商询价",
    ])
    
    gen.add_heading("5.5 图纸管理", level=2)
    gen.add_paragraph("路径：PLM产品研发 → 图纸管理")
    gen.add_bullet_list([
        "上传/下载图纸文件",
        "版本控制",
        "按项目/物料关联",
    ])
    gen.add_image("drawings.png", "图纸管理")
    
    gen.add_page_break()
    
    # ========== 第6章：ERP销售执行 ==========
    gen.add_heading("6. ERP销售执行", level=1)
    
    gen.add_heading("6.1 销售订单", level=2)
    gen.add_paragraph("路径：ERP销售执行 → 销售订单")
    gen.add_paragraph("管理正式销售订单。")
    gen.add_image("sales_order.png", "销售订单列表")
    
    gen.add_heading("6.1.1 新建销售订单", level=3)
    gen.add_numbered_list([
        "点击「新建订单」或从报价单转换",
        "填写订单信息：客户（必填）、客户订单号、交货日期、付款条款",
        "添加产品明细",
        "保存/提交审批",
    ])
    
    gen.add_heading("6.1.2 批量导入", level=3)
    gen.add_numbered_list([
        "下载导入模板",
        "按模板格式填写数据",
        "主表填写订单信息，明细表填写产品",
        "明细表需填写「订单号」以关联主表",
        "上传导入",
    ])
    gen.add_warning("客户名称必须与系统中已有客户匹配，否则导入失败。")
    
    gen.add_heading("6.2 发货管理", level=2)
    gen.add_paragraph("路径：ERP销售执行 → 发货管理")
    gen.add_paragraph("管理销售订单的发货：")
    gen.add_numbered_list([
        "从销售订单生成发货单",
        "选择发货数量",
        "填写物流信息",
        "确认发货",
    ])
    gen.add_image("delivery.png", "发货管理")
    
    gen.add_heading("6.3 应收账款", level=2)
    gen.add_paragraph("路径：ERP销售执行 → 应收账款")
    gen.add_bullet_list([
        "查看应收明细",
        "账龄分析",
        "催款管理",
    ])
    
    gen.add_page_break()
    
    # ========== 第7章：ERP采购管理 ==========
    gen.add_heading("7. ERP采购管理", level=1)
    
    gen.add_heading("7.1 采购申请", level=2)
    gen.add_paragraph("路径：ERP采购管理 → 采购申请")
    gen.add_paragraph("管理采购申请的全流程。")
    gen.add_image("purchase_request.png", "采购申请页面")
    
    gen.add_heading("7.1.1 页面布局", level=3)
    gen.add_paragraph("采购申请页面分为两部分：")
    gen.add_bullet_list([
        "物料需求清单（上方）：显示所有项目中待采购的物料（状态为'未下单'）",
        "采购申请列表（下方）：已创建的采购申请记录",
    ])
    
    gen.add_heading("7.1.2 创建采购申请", level=3)
    gen.add_paragraph("方式一：手动创建")
    gen.add_numbered_list([
        "点击「创建申请」",
        "选择项目",
        "从物料下拉框选择物料（仅显示该项目未采购的物料）",
        "填写数量、期望交期、供应商等",
        "保存",
    ])
    gen.add_paragraph("方式二：导入创建")
    gen.add_numbered_list([
        "点击「导入/导出」→「下载导入模板」",
        "填写模板（项目编号、物料编码、数量等）",
        "点击「导入/导出」→「导入采购申请」",
        "选择项目，上传文件",
    ])
    
    gen.add_heading("7.1.3 采购申请状态", level=3)
    gen.add_table(
        ["状态", "说明"],
        [
            ["草稿", "新建状态，可编辑"],
            ["待审批", "已提交，等待审批"],
            ["已批准", "审批通过"],
            ["已驳回", "审批驳回"],
            ["已转订单", "已转为采购订单"],
            ["已取消", "已取消"],
        ]
    )
    
    gen.add_heading("7.2 采购订单", level=2)
    gen.add_paragraph("路径：ERP采购管理 → 采购订单")
    gen.add_image("purchase_order.png", "采购订单")
    
    gen.add_paragraph("采购订单状态：")
    gen.add_table(
        ["状态", "说明"],
        [
            ["草稿", "新建状态"],
            ["已确认", "确认后通知供应商，BOM状态变为'已下单'"],
            ["执行中", "供应商已开始备货/发货"],
            ["已收货", "已完成到货验收"],
            ["已完成", "订单完结"],
            ["已取消", "订单已取消"],
        ]
    )
    
    gen.add_heading("7.3 到货质检", level=2)
    gen.add_paragraph("路径：ERP采购管理 → 到货质检")
    gen.add_paragraph("收货流程：")
    gen.add_numbered_list([
        "选择采购订单",
        "录入到货数量",
        "质量检验",
        "确认入库",
    ])
    gen.add_image("goods_receipt.png", "到货质检")
    
    gen.add_page_break()
    
    # ========== 第8章：ERP供应商管理 ==========
    gen.add_heading("8. ERP供应商管理", level=1)
    
    gen.add_heading("8.1 供应商档案", level=2)
    gen.add_paragraph("路径：ERP供应商管理 → 供应商档案")
    gen.add_table(
        ["字段", "说明"],
        [
            ["供应商编码", "唯一标识"],
            ["供应商名称", "公司全称"],
            ["联系人/电话", "主要联系人信息"],
            ["银行账户", "付款信息"],
            ["主营产品", "主要供应的产品类别"],
        ]
    )
    gen.add_image("supplier_list.png", "供应商列表")
    
    gen.add_heading("8.2 供应商评价", level=2)
    gen.add_paragraph("评价维度：")
    gen.add_bullet_list([
        "价格竞争力",
        "交货准时率",
        "质量合格率",
        "服务响应速度",
    ])
    
    gen.add_page_break()
    
    # ========== 第9章：ERP库存管理 ==========
    gen.add_heading("9. ERP库存管理", level=1)
    
    gen.add_heading("9.1 库存查询", level=2)
    gen.add_paragraph("路径：ERP库存管理 → 库存查询")
    gen.add_paragraph("查询维度：")
    gen.add_bullet_list([
        "按仓库",
        "按物料",
        "按批次",
        "按库位",
    ])
    gen.add_image("inventory.png", "库存查询")
    
    gen.add_heading("9.2 库存流水", level=2)
    gen.add_paragraph("路径：ERP库存管理 → 库存流水")
    gen.add_paragraph("流水类型：")
    gen.add_bullet_list([
        "采购入库",
        "销售出库",
        "生产领料",
        "生产退料",
        "调拨",
        "盘点调整",
    ])
    
    gen.add_heading("9.3 库存盘点", level=2)
    gen.add_paragraph("路径：ERP库存管理 → 库存盘点")
    gen.add_paragraph("盘点流程：")
    gen.add_numbered_list([
        "创建盘点单",
        "选择盘点范围",
        "录入实际数量",
        "生成差异报告",
        "审批调整",
    ])
    
    gen.add_page_break()
    
    # ========== 第10章：ERP财务管理 ==========
    gen.add_heading("10. ERP财务管理", level=1)
    
    gen.add_heading("10.1 费用报销", level=2)
    gen.add_paragraph("路径：ERP财务管理 → 费用报销")
    gen.add_paragraph("报销流程：")
    gen.add_numbered_list([
        "填写报销单",
        "上传发票/凭证",
        "提交审批",
        "财务审核",
        "付款",
    ])
    gen.add_image("expense.png", "费用报销")
    
    gen.add_heading("10.2 项目成本", level=2)
    gen.add_paragraph("路径：ERP财务管理 → 项目成本")
    gen.add_paragraph("成本构成：")
    gen.add_bullet_list([
        "材料成本（BOM物料）",
        "人工成本（工时×费率）",
        "费用成本（报销费用）",
        "外协成本",
    ])
    
    gen.add_page_break()
    
    # ========== 第11章：MES生产执行 ==========
    gen.add_heading("11. MES生产执行", level=1)
    
    gen.add_heading("11.1 生产计划", level=2)
    gen.add_paragraph("路径：MES生产执行 → 生产计划")
    gen.add_image("production_plan.png", "生产计划")
    
    gen.add_heading("11.2 工单派工", level=2)
    gen.add_paragraph("路径：MES生产执行 → 工单派工")
    gen.add_paragraph("分配生产工单给工人。")
    
    gen.add_heading("11.3 生产领料", level=2)
    gen.add_paragraph("路径：MES生产执行 → 生产领料")
    gen.add_numbered_list([
        "创建领料申请",
        "选择项目/工单",
        "选择物料和数量",
        "仓库确认出库",
    ])
    
    gen.add_heading("11.4 质量检验", level=2)
    gen.add_paragraph("路径：MES生产执行 → 质量检验")
    gen.add_paragraph("管理生产过程质量检验。")
    
    gen.add_page_break()
    
    # ========== 第12章：MES设备管理 ==========
    gen.add_heading("12. MES设备管理", level=1)
    
    gen.add_heading("12.1 设备台账", level=2)
    gen.add_paragraph("路径：MES设备管理 → 设备台账")
    gen.add_table(
        ["字段", "说明"],
        [
            ["设备编号", "唯一标识"],
            ["设备名称", "设备名称"],
            ["规格型号", "型号规格"],
            ["所属部门", "使用部门"],
            ["设备状态", "正常/维修/报废"],
        ]
    )
    gen.add_image("equipment.png", "设备台账")
    
    gen.add_heading("12.2 设备点检", level=2)
    gen.add_paragraph("路径：MES设备管理 → 设备点检")
    gen.add_paragraph("执行设备日常点检。")
    
    gen.add_heading("12.3 OEE分析", level=2)
    gen.add_paragraph("路径：MES设备管理 → OEE分析")
    gen.add_paragraph("设备综合效率分析（Overall Equipment Effectiveness）：")
    gen.add_paragraph("OEE = 可用率 × 性能率 × 质量率", bold=True)
    
    gen.add_page_break()
    
    # ========== 第13章：OA协同办公 ==========
    gen.add_heading("13. OA协同办公", level=1)
    
    gen.add_heading("13.1 待办审批", level=2)
    gen.add_paragraph("路径：OA协同办公 → 待办审批")
    gen.add_paragraph("处理待审批事项，可审批类型：")
    gen.add_bullet_list([
        "采购申请",
        "费用报销",
        "请假申请",
        "用车申请",
        "等",
    ])
    gen.add_image("workflow.png", "待办审批")
    
    gen.add_heading("13.2 我的提交", level=2)
    gen.add_paragraph("路径：OA协同办公 → 我的提交")
    gen.add_paragraph("查看我提交的审批记录及其状态。")
    
    gen.add_heading("13.3 请假申请", level=2)
    gen.add_paragraph("路径：OA协同办公 → 请假申请")
    gen.add_paragraph("请假类型：")
    gen.add_bullet_list([
        "事假",
        "病假",
        "年假",
        "调休",
        "其他",
    ])
    
    gen.add_heading("13.4 会议管理", level=2)
    gen.add_paragraph("路径：OA协同办公 → 会议管理")
    gen.add_paragraph("预约和管理会议。")
    
    gen.add_page_break()
    
    # ========== 第14章：系统管理 ==========
    gen.add_heading("14. 系统管理", level=1)
    
    gen.add_heading("14.1 用户管理", level=2)
    gen.add_paragraph("路径：系统管理 → 用户管理")
    gen.add_bullet_list([
        "创建/编辑/禁用用户",
        "分配角色",
        "重置密码",
    ])
    gen.add_image("user_list.png", "用户管理")
    
    gen.add_heading("14.2 角色管理", level=2)
    gen.add_paragraph("路径：系统管理 → 角色管理")
    gen.add_paragraph("权限类型：")
    gen.add_table(
        ["权限类型", "说明"],
        [
            ["菜单权限", "控制可访问的菜单"],
            ["操作权限", "控制可执行的操作（增删改查）"],
            ["数据权限", "控制可见的数据范围（本人/本部门/全部）"],
        ]
    )
    gen.add_image("role_list.png", "角色管理")
    
    gen.add_heading("14.3 编码规则", level=2)
    gen.add_paragraph("路径：系统管理 → 编码规则")
    gen.add_paragraph("配置各类单据的编号规则，示例：")
    gen.add_bullet_list([
        "销售订单：SO-YYYYMMDD-####",
        "采购订单：PO-YYYYMMDD-####",
        "项目编号：PRJ-YYYYMMDD-####",
    ])
    
    gen.add_page_break()
    
    # ========== 第15章：报表分析 ==========
    gen.add_heading("15. 报表分析", level=1)
    
    gen.add_paragraph("系统提供丰富的报表分析功能：")
    gen.add_table(
        ["报表", "说明"],
        [
            ["项目利润分析", "分析项目的收入、成本和利润"],
            ["库存周转率", "分析库存周转效率"],
            ["账龄分析", "分析应收/应付账款的账龄分布"],
            ["采购价格趋势", "分析物料采购价格变化趋势"],
            ["现金流预测", "预测未来现金流入/流出"],
            ["呆滞物料分析", "分析长期不动的呆滞库存"],
            ["工时统计", "统计员工工时"],
            ["项目成本分析", "分析项目成本构成"],
        ]
    )
    gen.add_image("reports.png", "报表分析")
    
    gen.add_page_break()
    
    # ========== 第16章：常见问题 ==========
    gen.add_heading("16. 常见问题FAQ", level=1)
    
    gen.add_heading("Q1: 忘记密码怎么办？", level=2)
    gen.add_paragraph("联系系统管理员重置密码。")
    
    gen.add_heading("Q2: 如何批量导入数据？", level=2)
    gen.add_numbered_list([
        "进入相应模块（如客户管理、销售订单）",
        "点击「导入」或「导入/导出」按钮",
        "下载导入模板",
        "按模板格式填写数据",
        "上传文件完成导入",
    ])
    
    gen.add_heading("Q3: 导入失败怎么处理？", level=2)
    gen.add_bullet_list([
        "系统会显示详细的错误信息",
        "根据错误提示修改数据",
        "常见问题：客户名称不匹配、必填字段为空、数据格式错误",
    ])
    
    gen.add_heading("Q4: 如何查看审批进度？", level=2)
    gen.add_paragraph("进入「OA协同办公 → 我的提交」查看已提交的审批及其状态。")
    
    gen.add_heading("Q5: 采购申请导入后物料需求清单没更新？", level=2)
    gen.add_bullet_list([
        "导入后系统会自动刷新列表",
        "如未更新，请手动刷新页面",
    ])
    
    gen.add_heading("Q6: BOM采购状态不更新怎么办？", level=2)
    gen.add_paragraph("检查采购申请/订单的操作是否完成：")
    gen.add_bullet_list([
        "创建采购申请 → BOM状态变为'申请中'",
        "批准采购申请 → BOM状态变为'已批准'",
        "确认采购订单 → BOM状态变为'已下单'",
    ])
    
    gen.add_heading("Q7: 销售订单导入时明细没关联上怎么办？", level=2)
    gen.add_paragraph("确保明细表中填写了对应的「订单号」字段，用于关联主表的订单。")
    
    gen.add_heading("Q8: 已关联项目的销售订单还能选择吗？", level=2)
    gen.add_paragraph("不能。已关联项目的订单会从可选列表中移除。如需重新关联，需先取消或删除原项目。")
    
    gen.add_page_break()
    
    # 附录
    gen.add_heading("附录", level=1)
    
    gen.add_heading("A. 快捷键", level=2)
    gen.add_table(
        ["快捷键", "功能"],
        [
            ["Ctrl + S", "保存"],
            ["Ctrl + Enter", "提交"],
            ["Esc", "关闭弹窗"],
        ]
    )
    
    gen.add_heading("B. 截图文件命名规范", level=2)
    gen.add_paragraph("如需添加系统截图，请将截图放入 screenshots/ 目录，按以下命名：")
    gen.add_table(
        ["文件名", "说明"],
        [
            ["login.png", "登录页面"],
            ["dashboard.png", "工作台"],
            ["customer_list.png", "客户列表"],
            ["customer_form.png", "客户表单"],
            ["project_list.png", "项目列表"],
            ["project_form.png", "项目表单"],
            ["bom_list.png", "BOM列表"],
            ["purchase_request.png", "采购申请"],
            ["purchase_order.png", "采购订单"],
            ["sales_order.png", "销售订单"],
            ["inventory.png", "库存查询"],
            ["workflow.png", "审批流程"],
        ]
    )
    
    # 保存
    gen.save("ERP系统用户操作手册.docx")


if __name__ == "__main__":
    create_manual()
